"""Resume generation service — produces a professional DOCX from structured content."""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.llm.schemas import Recommendation, ResumeContent, StructuredJD, StructuredResume
from src.llm.service import get_llm_service

if TYPE_CHECKING:
    from src.generator.schemas import GenerateRequest


# ── DOCX rendering ──────────────────────────────────────────────────────────


def _set_run_font(run, *, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:  # type: ignore[no-untyped-def]
    """Apply common font properties to a python-docx Run."""
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a styled section heading with a subtle bottom border."""
    para = doc.add_paragraph()
    para.space_before = Pt(12)
    para.space_after = Pt(4)
    run = para.add_run(text.upper())
    _set_run_font(run, size=12, bold=True, color=RGBColor(0x1A, 0x56, 0xDB))
    # Add a thin bottom-border via paragraph shading workaround
    pf = para.paragraph_format
    pf.space_after = Pt(2)


def generate_docx(content: ResumeContent) -> bytes:
    """Render a ``ResumeContent`` object into a professional DOCX file and return bytes."""
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── Header: name ─────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(content.name)
    _set_run_font(name_run, size=22, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))

    # ── Contact info ─────────────────────────────────────────────────────
    contact_parts: list[str] = []
    if content.email:
        contact_parts.append(content.email)
    if content.phone:
        contact_parts.append(content.phone)
    if content.location:
        contact_parts.append(content.location)
    if content.linkedin:
        contact_parts.append(content.linkedin)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        _set_run_font(contact_run, size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ── Summary ──────────────────────────────────────────────────────────
    if content.summary:
        _add_section_heading(doc, "Professional Summary")
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(content.summary)
        _set_run_font(summary_run, size=11)

    # ── Skills ───────────────────────────────────────────────────────────
    if content.skills:
        _add_section_heading(doc, "Skills")
        skills_para = doc.add_paragraph()
        skills_run = skills_para.add_run("  •  ".join(content.skills))
        _set_run_font(skills_run, size=11)

    # ── Experience ───────────────────────────────────────────────────────
    if content.experience:
        _add_section_heading(doc, "Professional Experience")
        for exp in content.experience:
            # Title + Company
            title_para = doc.add_paragraph()
            title_para.space_before = Pt(6)
            title_para.space_after = Pt(0)
            title_run = title_para.add_run(exp.title)
            _set_run_font(title_run, size=11, bold=True)
            sep_run = title_para.add_run(f"  —  {exp.company}")
            _set_run_font(sep_run, size=11)

            # Duration
            dur_para = doc.add_paragraph()
            dur_para.space_before = Pt(0)
            dur_para.space_after = Pt(2)
            dur_run = dur_para.add_run(exp.duration)
            _set_run_font(dur_run, size=10, color=RGBColor(0x77, 0x77, 0x77))

            # Description bullets
            for line in exp.description.split("\n"):
                line = line.strip().lstrip("•-–").strip()
                if line:
                    bullet = doc.add_paragraph(style="List Bullet")
                    bullet.space_before = Pt(0)
                    bullet.space_after = Pt(1)
                    bullet_run = bullet.add_run(line)
                    _set_run_font(bullet_run, size=10)

            # Technologies
            if exp.technologies:
                tech_para = doc.add_paragraph()
                tech_para.space_before = Pt(2)
                tech_label = tech_para.add_run("Technologies: ")
                _set_run_font(tech_label, size=10, bold=True, color=RGBColor(0x55, 0x55, 0x55))
                tech_val = tech_para.add_run(", ".join(exp.technologies))
                _set_run_font(tech_val, size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ── Education ────────────────────────────────────────────────────────
    if content.education:
        _add_section_heading(doc, "Education")
        for edu in content.education:
            edu_para = doc.add_paragraph()
            edu_para.space_before = Pt(4)
            edu_para.space_after = Pt(1)
            deg_run = edu_para.add_run(edu.degree)
            _set_run_font(deg_run, size=11, bold=True)
            inst_run = edu_para.add_run(f"  —  {edu.institution}")
            _set_run_font(inst_run, size=11)
            if edu.year:
                year_run = edu_para.add_run(f"  ({edu.year})")
                _set_run_font(year_run, size=10, color=RGBColor(0x77, 0x77, 0x77))
            if edu.gpa:
                gpa_para = doc.add_paragraph()
                gpa_para.space_before = Pt(0)
                gpa_run = gpa_para.add_run(f"GPA: {edu.gpa}")
                _set_run_font(gpa_run, size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ── Certifications ───────────────────────────────────────────────────
    if content.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in content.certifications:
            cert_para = doc.add_paragraph(style="List Bullet")
            cert_para.space_before = Pt(1)
            cert_para.space_after = Pt(1)
            cert_run = cert_para.add_run(cert)
            _set_run_font(cert_run, size=11)

    # ── Serialise to bytes ───────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ── High-level generation pipeline ───────────────────────────────────────────


async def generate_resume(request: "GenerateRequest") -> bytes:
    """Produce a DOCX resume from the request.

    Modes:
    - ``enhanced`` / ``accept_all``: incorporate accepted recommendations.
    - ``overhaul``: complete LLM rewrite targeting the JD.
    """
    llm_service = get_llm_service()

    azure_kwargs = {
        "azure_endpoint": request.azure_endpoint,
        "azure_deployment": request.azure_deployment,
        "azure_api_version": request.azure_api_version,
    }

    # If we don't already have structured data, parse it now
    structured_resume = request.structured_resume
    if structured_resume is None:
        structured_resume = await llm_service.analyze_resume(
            request.resume_text, request.provider, request.api_key, **azure_kwargs
        )

    structured_jd = request.structured_jd
    if structured_jd is None:
        structured_jd = await llm_service.analyze_jd(
            request.jd_text, request.provider, request.api_key, **azure_kwargs
        )

    # Convert accepted recommendations to the LLM schema objects
    from src.llm.schemas import Recommendation as LLMRec

    recs = [
        LLMRec(
            id=ar.id,
            category=ar.category,
            title="",
            original_text=None,
            suggested_text=ar.suggested_text,
            reasoning="",
            priority="high",
            action=ar.action,
        )
        for ar in request.accepted_recommendations
    ]

    content: ResumeContent = await llm_service.generate_resume_content(
        resume=structured_resume,
        recommendations=recs,
        jd=structured_jd,
        mode=request.mode,
        provider=request.provider,
        api_key=request.api_key,
        **azure_kwargs
    )

    return generate_docx(content)
